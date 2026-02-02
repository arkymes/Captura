import os
import re
import sys
import json
from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any
from datetime import datetime
import textwrap

import cv2  # type: ignore
from bs4 import BeautifulSoup, NavigableString, Tag  # type: ignore
from markdown import markdown  # type: ignore
import unicodedata
import shutil
from html import escape
import requests  # type: ignore

def clean_text(text: str) -> str:
    """Limpa texto removendo caracteres de controle e normalizando Unicode."""
    # Remove caracteres nulos e de controle (exceto \n \t \r)
    text = ''.join(c for c in text if ord(c) >= 32 or c in '\n\t\r')
    # Normaliza Unicode
    text = unicodedata.normalize('NFC', text)
    # Remove surrogates e caracteres inválidos
    text = text.encode('utf-8', errors='replace').decode('utf-8')
    return text

from docx import Document  # type: ignore
from docx.shared import Pt, Inches, Mm  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL  # type: ignore
from docx.oxml import OxmlElement  # type: ignore
from docx.oxml.ns import qn  # type: ignore


ROOT = Path(__file__).resolve().parent
# WORKDIR agora é o diretório raiz do projeto (pai de Captura)
WORKDIR = ROOT.parent
DOC_MD = WORKDIR / "doc.md"
# Permite definir o vídeo via variável de ambiente INPUT_VIDEO_PATH; caso contrário usa um padrão no diretório de trabalho
_VIDEO_ENV = os.environ.get("INPUT_VIDEO_PATH")
if _VIDEO_ENV:
    VIDEO_FILE = Path(_VIDEO_ENV)
else:
    # Padrão esperado quando rodado pelo app: arquivo salvo como "input_video.ext" dentro de GeraçãoDOc
    VIDEO_FILE = next((p for p in WORKDIR.glob("input_video.*") if p.is_file()), WORKDIR / "input_video.mp4")
OUT_DIR = Path(os.environ.get("OUTPUT_DIR", str(WORKDIR / "docs")))
ASSETS_BASE = OUT_DIR / "assets"
ASSETS_DIR = ASSETS_BASE / "prints"
DIAGRAMS_DIR = ASSETS_BASE / "diagrams"
INDEX_DOCX = OUT_DIR / "index.docx"
META_FILES = [
    WORKDIR / "doc_meta.json",
]
MODEL_ASSETS_DIR = WORKDIR / "Documento_exemplo" / "1760722885_it.fin.xx.-conciliao-contas-a-receber"

MAX_DIAGRAM_WIDTH_INCHES = 5.2
MAX_DIAGRAM_HEIGHT_INCHES = 8.5
PRINT_IMAGE_SCALE = 0.75

# Área útil do vídeo (remove barras pretas e overlays do Teams)
CROP_LEFT_X = 0
CROP_RIGHT_X = 1675
CROP_TOP_Y = 70
CROP_BOTTOM_Y = 1010

PRINT_PATTERN = re.compile(
    r"(?P<prefix>^[ \t]*[-*]\s*)?"  # marcador de lista opcional
    r"\[PRINT DO V[ÍI]DEO\s*-\s*(?P<ts>(?:\d{1,2}:)?\d{1,2}:\d{2})\s*:\s*(?P<desc>.*?)\]"
    r"(?:\s*(?:\n\s*(?:[-*]\s*)?)?\{(?P<coords>[^{}]+)\})?",
    re.IGNORECASE | re.MULTILINE,
)
# Padrão corrigido para capturar blocos mermaid de forma mais robusta
MERMAID_BLOCK_PATTERN = re.compile(
    r"```mermaid[ \t]*\r?\n([\s\S]*?)\r?\n[ \t]*```",
    re.IGNORECASE,
)
PRINT_TOKEN_PATTERN = re.compile(r"PRINT_SLOT_(\d+)_TOKEN")


@dataclass
class PrintOccurrence:
    token: str
    timestamp: str
    description: str
    image_path: Path
    coords: Optional[Tuple[int, int]]


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_BASE.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    DIAGRAMS_DIR.mkdir(parents=True, exist_ok=True)


def _deduce_defaults_from_title(title_text: str) -> dict:
    title_text = (title_text or "").strip()
    doc_code = ""
    doc_title = title_text or "Documento"
    if title_text:
        parts = re.split(r"\s*[-–—]\s*", title_text, maxsplit=1)
        if len(parts) == 2:
            doc_code = parts[0].strip()
            doc_title = parts[1].strip() or doc_title
    defaults = {
        "doc_type": "INSTRUÇÃO DE TRABALHO",
        "doc_code": doc_code or "—",
        "doc_title": doc_title or "Documento",
        "doc_issue": datetime.now().strftime('%d/%m/%Y'),
        "doc_revision": "1.0",
        "elaboracao": "A definir",
        "aprovacao": "A definir",
        "empresa": "Empresa",
    }
    return defaults


_META_KEY_ALIASES = {
    "doc_type": ["doc_type", "tipo", "tipo_documento"],
    "doc_code": ["doc_code", "codigo", "código"],
    "doc_title": ["doc_title", "titulo", "título"],
    "doc_issue": ["doc_issue", "data", "data_emissao", "data_emissão"],
    "doc_revision": ["doc_revision", "revisao", "revisão"],
    "elaboracao": ["elaboracao", "elaboração", "elaboracao_revisao"],
    "aprovacao": ["aprovacao", "aprovação"],
    "empresa": ["empresa", "companhia"],
}


def _normalize_metadata_keys(raw: dict) -> dict:
    if not isinstance(raw, dict):
        return {}
    normalized: dict = {}
    lower_map = {k: v for k, v in raw.items()}
    for canonical, options in _META_KEY_ALIASES.items():
        for key in options:
            if key in lower_map:
                normalized[canonical] = lower_map[key]
                break
    return normalized


def load_metadata(title_text: str) -> dict:
    metadata = _deduce_defaults_from_title(title_text)
    for candidate in META_FILES:
        if not candidate.exists():
            continue
        try:
            data = json.loads(candidate.read_text(encoding='utf-8'))
        except Exception as exc:  # noqa: BLE001
            print(f"[AVISO] Falha ao ler metadata em {candidate}: {exc}")
            continue
        normalized = _normalize_metadata_keys(data)
        for key, value in normalized.items():
            if value is not None:
                metadata[key] = str(value)
    metadata.setdefault("empresa", "Empresa")
    return metadata


def find_logo(layout_assets_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Busca o logo em ordem de prioridade:
    1. Layout assets configurado pelo usuário
    2. Arquivos padrão no workdir
    3. Assets do modelo de exemplo
    """
    # Prioridade 1: Layout assets configurado
    if layout_assets_dir and layout_assets_dir.exists():
        for file in layout_assets_dir.glob("logo_*"):
            if file.is_file():
                dest = ASSETS_BASE / file.name
                try:
                    shutil.copyfile(file, dest)
                    return dest
                except Exception:
                    return file
    
    # Prioridade 2: Arquivos padrão
    candidates = [
        WORKDIR / "Logo_empresa.png",
        WORKDIR / "logo_empresa.png",
        WORKDIR / "logo.png",
    ]
    for candidate in candidates:
        if candidate.exists():
            dest = ASSETS_BASE / candidate.name
            try:
                shutil.copyfile(candidate, dest)
                return dest
            except Exception:
                return candidate
    
    return None


def find_model_separator(layout_assets_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Busca o separador em ordem de prioridade:
    1. Layout assets configurado pelo usuário
    2. Assets do modelo de exemplo
    """
    # Prioridade 1: Layout assets configurado
    if layout_assets_dir and layout_assets_dir.exists():
        for file in layout_assets_dir.glob("separator_*"):
            if file.is_file():
                dest = ASSETS_BASE / file.name
                try:
                    shutil.copyfile(file, dest)
                    return dest
                except Exception:
                    return file
    
    # Prioridade 2: Modelo de exemplo
    candidate = MODEL_ASSETS_DIR / "1760722885_it.fin.xx.-conciliao-contas-a-receber-2.png"
    if candidate.exists():
        dest = ASSETS_BASE / candidate.name
        try:
            shutil.copyfile(candidate, dest)
            return dest
        except Exception:
            return candidate
    
    return None


def find_model_footer_banner(layout_assets_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Busca o banner do rodapé em ordem de prioridade:
    1. Layout assets configurado pelo usuário
    2. Assets do modelo de exemplo
    """
    # Prioridade 1: Layout assets configurado
    if layout_assets_dir and layout_assets_dir.exists():
        for file in layout_assets_dir.glob("footer_*"):
            if file.is_file():
                dest = ASSETS_BASE / file.name
                try:
                    shutil.copyfile(file, dest)
                    return dest
                except Exception:
                    return file
    
    # Prioridade 2: Modelo de exemplo
    candidate = MODEL_ASSETS_DIR / "1760722885_it.fin.xx.-conciliao-contas-a-receber-17.png"
    if candidate.exists():
        dest = ASSETS_BASE / candidate.name
        try:
            shutil.copyfile(candidate, dest)
            return dest
        except Exception:
            return candidate
    
    return None


def timestamp_to_seconds(mm: str, ss: str) -> float:
    m = int(mm)
    s = int(ss)
    return m * 60 + s


def parse_timestamp_to_seconds(ts: str) -> float:
    ts = (ts or "").strip()
    if not ts:
        return 0.0
    parts = ts.split(":")
    try:
        if len(parts) == 3:
            h, m, s = [int(p) for p in parts]
            return h * 3600 + m * 60 + s
        if len(parts) == 2:
            m, s = [int(p) for p in parts]
            return m * 60 + s
        if len(parts) == 1:
            return float(int(parts[0]))
    except Exception:
        return 0.0
    return 0.0


def _write_image_bgr(img, out_path: Path) -> bool:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    suffix = out_path.suffix.lower()
    if suffix not in {'.jpg', '.jpeg', '.png'}:
        suffix = '.jpg'

    if suffix == '.png':
        success, buf = cv2.imencode('.png', img)
    else:
        success, buf = cv2.imencode('.jpg', img, [int(cv2.IMWRITE_JPEG_QUALITY), 90])
    if not success:
        return False
    try:
        out_path.write_bytes(buf.tobytes())
        return True
    except Exception as exc:
        print(f"[ERRO] Falha ao salvar imagem: {out_path} - {exc}")
        return False


def _make_placeholder(ts_seconds: float, label: str) -> 'cv2.Mat':
    width, height = 1280, 720
    bg = (240, 243, 247)  # BGR light gray
    import numpy as np  # type: ignore
    img = np.zeros((height, width, 3), dtype='uint8')
    img[:] = bg

    cv2.rectangle(img, (0, 0), (width, 8), (27, 132, 204), -1)
    cv2.rectangle(img, (0, height - 8), (width, height), (205, 7, 11), -1)

    mm = int(ts_seconds // 60)
    ss = int(ts_seconds % 60)
    ts_str = f"{mm:02d}:{ss:02d}"
    title = f"PRINT DO VÍDEO - {ts_str}"
    subtitle = label[:70]

    cv2.putText(img, title, (40, 120), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (27, 132, 204), 2, cv2.LINE_AA)
    cv2.putText(img, subtitle, (40, 200), cv2.FONT_HERSHEY_SIMPLEX, 0.95, (80, 80, 80), 1, cv2.LINE_AA)
    note = "Frame não capturado - imagem ilustrativa"
    cv2.putText(img, note, (40, height - 25), cv2.FONT_HERSHEY_SIMPLEX, 0.75, (120, 120, 120), 1, cv2.LINE_AA)
    return img


def extract_frame(video_path: Path, seconds: float, out_path: Path, coords: Optional[Tuple[int, int]] = None) -> bool:
    cap = cv2.VideoCapture(str(video_path))
    if not cap.isOpened():
        placeholder = _make_placeholder(seconds, out_path.stem.replace('_', ' '))
        return _write_image_bgr(placeholder, out_path)

    fps = cap.get(cv2.CAP_PROP_FPS) or 30.0
    total_frames = cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0
    target = int(seconds * fps)
    if total_frames and target >= total_frames:
        target = max(int(total_frames) - 2, 0)
    cap.set(cv2.CAP_PROP_POS_FRAMES, max(target - 1, 0))
    ok, frame = cap.read()
    if not ok or frame is None:
        cap.release()
        placeholder = _make_placeholder(seconds, out_path.stem.replace('_', ' '))
        return _write_image_bgr(placeholder, out_path)

    # Recorte para remover barras e overlays (topo, direita e rodapé)
    h, w = frame.shape[:2]
    crop_left = max(0, min(CROP_LEFT_X, max(w - 1, 0)))
    crop_right = w if CROP_RIGHT_X <= 0 else max(crop_left + 1, min(CROP_RIGHT_X, w))
    crop_top = max(0, min(CROP_TOP_Y, max(h - 1, 0)))
    crop_bottom = max(crop_top + 1, min(CROP_BOTTOM_Y, h))
    if crop_bottom <= crop_top:
        crop_bottom = h
        crop_top = 0
    if crop_bottom > crop_top:
        frame = frame[crop_top:crop_bottom, crop_left:crop_right]
        h, w = frame.shape[:2]

    max_width = 1400
    if w > max_width:
        scale = max_width / w
        frame = cv2.resize(frame, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)

    ok = _write_image_bgr(frame, out_path)
    cap.release()
    if not ok:
        placeholder = _make_placeholder(seconds, out_path.stem.replace('_', ' '))
        return _write_image_bgr(placeholder, out_path)
    return ok


_COORD_EXTRACT_PATTERN = re.compile(r"\{([^{}]+)\}\s*$")
_COORD_PAIR_PATTERN = re.compile(r"(x|y)\s*=\s*(-?\d+)", re.IGNORECASE)


def replace_print_placeholders(md_text: str) -> Tuple[str, List[PrintOccurrence]]:
    occurrences: List[PrintOccurrence] = []

    def repl(match: re.Match) -> str:
        ts_str = (match.group('ts') or '').strip()
        raw_desc = (match.group('desc') or '').strip()
        prefix = match.group('prefix') or ''

        coords: Optional[Tuple[int, int]] = None
        coords_text = match.group('coords')
        if coords_text:
            pairs = {k.lower(): int(v) for k, v in _COORD_PAIR_PATTERN.findall(coords_text)}
            if 'x' in pairs and 'y' in pairs:
                coords = (pairs['x'], pairs['y'])
        else:
            coord_match = _COORD_EXTRACT_PATTERN.search(raw_desc)
            if coord_match:
                coords_text = coord_match.group(1)
                raw_desc = raw_desc[:coord_match.start()].strip()
                pairs = {k.lower(): int(v) for k, v in _COORD_PAIR_PATTERN.findall(coords_text)}
                if 'x' in pairs and 'y' in pairs:
                    coords = (pairs['x'], pairs['y'])

        desc = raw_desc
        ts_for_file = ts_str.replace(":", "-")
        safe_desc = re.sub(r"[^a-zA-Z0-9_-]+", "-", desc)[:80].strip("-") or "print"
        filename = f"frame_{ts_for_file}_{safe_desc}.jpg".lower()
        img_path = ASSETS_DIR / filename
        token = f"PRINT_SLOT_{len(occurrences) + 1}_TOKEN"
        occurrences.append(PrintOccurrence(
            token=token,
            timestamp=ts_str,
            description=desc,
            image_path=img_path,
            coords=coords,
        ))
        replacement = token
        if prefix:
            replacement = f"{prefix}{token}"
        return replacement

    new_md = PRINT_PATTERN.sub(repl, md_text)
    return new_md, occurrences


def _create_diagram_placeholder(output_path: Path) -> bool:
    width, height = 1280, 720
    try:
        import numpy as np  # type: ignore
    except Exception as exc:  # noqa: BLE001
        print(f"[AVISO] Falha ao importar numpy para placeholder de diagrama: {exc}")
        return False

    img = np.full((height, width, 3), 255, dtype='uint8')
    cv2.rectangle(img, (40, 40), (width - 40, height - 40), (200, 200, 200), thickness=4)
    cv2.putText(img, 'Diagrama não renderizado', (60, height // 2 - 30), cv2.FONT_HERSHEY_SIMPLEX, 1.2, (50, 50, 50), 3, cv2.LINE_AA)
    cv2.putText(img, 'Verifique a conexão ou gere manualmente', (60, height // 2 + 40), cv2.FONT_HERSHEY_SIMPLEX, 0.8, (90, 90, 90), 2, cv2.LINE_AA)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    max_export_width = 1200
    max_export_height = 720
    current_height, current_width = img.shape[:2]
    export_scale = min(1.0, max_export_width / max(current_width, 1), max_export_height / max(current_height, 1))
    if export_scale < 1.0:
        new_size = (max(1, int(current_width * export_scale)), max(1, int(current_height * export_scale)))
        img = cv2.resize(img, new_size, interpolation=cv2.INTER_AREA)

    return _write_image_bgr(img, output_path)


def _build_print_figure(soup: BeautifulSoup, occurrence: PrintOccurrence) -> Tag:
    figure = soup.new_tag("figure", attrs={"class": "video-print"})
    img = soup.new_tag("img", src=f"assets/prints/{occurrence.image_path.name}", alt=occurrence.description, loading="lazy")
    figure.append(img)
    caption = soup.new_tag("figcaption")
    caption.string = occurrence.description
    figure.append(caption)
    return figure


def inject_print_figures(soup: BeautifulSoup, occurrences: List[PrintOccurrence]) -> None:
    if not occurrences:
        return

    token_map: Dict[str, PrintOccurrence] = {occ.token: occ for occ in occurrences}
    for text_node in list(soup.find_all(string=PRINT_TOKEN_PATTERN)):
        text_value = str(text_node)
        matches = list(PRINT_TOKEN_PATTERN.finditer(text_value))
        if not matches:
            continue

        fragments: List[Any] = []
        last = 0
        for match in matches:
            if match.start() > last:
                fragments.append(text_value[last:match.start()])
            token = match.group(0)
            occurrence = token_map.get(token)
            if occurrence is None:
                fragments.append(token)
            else:
                figure_tag = _build_print_figure(soup, occurrence)
                fragments.append(figure_tag)
            last = match.end()
        if last < len(text_value):
            fragments.append(text_value[last:])

        for fragment in fragments:
            if isinstance(fragment, str):
                if fragment:
                    text_node.insert_before(fragment)
            else:
                text_node.insert_before(fragment)
        text_node.extract()


def _render_mermaid_diagram(code: str, output_path: Path) -> bool:
    output_path.parent.mkdir(parents=True, exist_ok=True)
    url = "https://kroki.io/mermaid/png"
    headers = {'Accept': 'image/png', 'Content-Type': 'application/json; charset=utf-8'}
    try:
        response = requests.post(url, json={'diagram_source': code}, headers=headers, timeout=20)
        if response.status_code == 200:
            output_path.write_bytes(response.content)
            return True
        preview = response.text[:200] if response.content else ''
        if preview:
            print(f"[AVISO] Kroki retornou status {response.status_code}: {preview}")
        else:
            print(f"[AVISO] Kroki retornou status {response.status_code} para diagrama Mermaid.")
    except Exception as error:  # noqa: BLE001
        print(f"[AVISO] Falha ao renderizar diagrama Mermaid via Kroki: {error}")
    if _render_mermaid_locally(code, output_path):
        print("[OK] Diagrama Mermaid renderizado localmente.")
        return True
    print("[AVISO] Usando placeholder para diagrama Mermaid.")
    return _create_diagram_placeholder(output_path)


def replace_mermaid_blocks(md_text: str) -> Tuple[str, List[Path]]:
    generated_paths: List[Path] = []

    def repl(match: re.Match) -> str:
        code = match.group(1)  # Grupo 1 agora contém o código do diagrama
        if code is None:
            return ''
        diagram_code = code.strip()
        if not diagram_code:
            return ''
        diagram_index = len(generated_paths) + 1
        filename = f"diagram_{diagram_index:02d}.png"
        output_path = DIAGRAMS_DIR / filename
        success = _render_mermaid_diagram(diagram_code, output_path)
        if success:
            print(f"[OK] Diagrama Mermaid #{diagram_index} salvo em: {output_path}")
        generated_paths.append(output_path)

        caption = "Diagrama de fluxo do processo"
        escaped_caption = escape(caption)
        figure_html = (
            f'<figure class="process-diagram">'
            f'<img src="assets/diagrams/{filename}" alt="{escaped_caption}" loading="lazy"/>'
            f'<figcaption>{escaped_caption}</figcaption>'
            '</figure>'
        )
        return figure_html

    new_md = MERMAID_BLOCK_PATTERN.sub(repl, md_text)
    return new_md, generated_paths


def _parse_node_token(token: str) -> Tuple[str, str, str]:
    """
    Parseia um token de nó Mermaid e extrai: node_id, label, shape.
    
    Exemplos:
        - "Start([Início])" -> ('Start', 'Início', 'circle')
        - "Config[Carregar Configurações]" -> ('Config', 'Carregar Configurações', 'rect')
        - "Classify[Classificar Transações <br/>(Regex no Histórico)]" -> ('Classify', 'Classificar Transações (Regex no Histórico)', 'rect')
    """
    token = token.strip()
    if not token:
        return '', '', 'rect'

    # Lista de delimitadores ordenados do mais específico para o menos específico
    # O formato é: (start_delim, end_delim, shape_name)
    shapes = [
        ('([', '])', 'circle'),   # Stadium shape - usado para início/fim
        ('[[', ']]', 'rect'),     # Subroutine
        ('{{', '}}', 'circle'),   # Hexagon
        ('((', '))', 'circle'),   # Circle
        ('{', '}', 'diamond'),    # Diamond/Decision
        ('[', ']', 'rect'),       # Rectangle
        ('(', ')', 'round'),      # Rounded rectangle
    ]
    
    node_id = token
    label = token
    shape = 'rect'
    
    # Procura o primeiro delimitador de abertura que aparece no token
    # e encontra seu correspondente delimitador de fechamento no FINAL do token
    for start_delim, end_delim, shape_name in shapes:
        # Verifica se o token contém o delimitador de abertura
        start_idx = token.find(start_delim)
        if start_idx == -1:
            continue
            
        # Verifica se o token TERMINA com o delimitador de fechamento correspondente
        if not token.endswith(end_delim):
            continue
        
        # Extrai o node_id (parte antes do delimitador de abertura)
        node_id = token[:start_idx].strip()
        
        # Extrai o label (conteúdo entre os delimitadores)
        label_start = start_idx + len(start_delim)
        label_end = len(token) - len(end_delim)
        label = token[label_start:label_end].strip()
        
        shape = shape_name
        break
    else:
        # Se não encontrou delimitadores específicos, tenta separar por espaço
        parts = token.split()
        if parts:
            node_id = parts[0].strip()
            label = ' '.join(parts[1:]).strip() or node_id

    # Limpa tags HTML do label (como <br/>)
    label = re.sub(r'<br\s*/?>', ' ', label)
    # Remove espaços duplicados
    label = re.sub(r'\s+', ' ', label).strip()

    if not node_id:
        node_id = label
    if not label:
        label = node_id
    return node_id, label, shape


def _register_node(
    nodes: dict,
    subgraph_nodes: dict,
    standalone_nodes: List[str],
    node_id: str,
    label: str,
    shape: str,
    current_subgraph: Optional[str],
) -> None:
    if not node_id:
        return
    info = nodes.setdefault(node_id, {'label': label, 'shape': shape, 'subgraph': current_subgraph})
    if label:
        info['label'] = label
    if not info.get('shape'):
        info['shape'] = shape
    if info.get('subgraph') is None and current_subgraph is not None:
        info['subgraph'] = current_subgraph
    if current_subgraph:
        subgraph_nodes.setdefault(current_subgraph, [])
        if node_id not in subgraph_nodes[current_subgraph]:
            subgraph_nodes[current_subgraph].append(node_id)
    else:
        if node_id not in standalone_nodes:
            standalone_nodes.append(node_id)


def _parse_mermaid_structure(code: str) -> dict:
    nodes: dict = {}
    edges: List[dict] = []
    subgraph_nodes: dict = {}
    subgraph_order: List[str] = []
    standalone_nodes: List[str] = []
    current_subgraph: Optional[str] = None

    for raw_line in code.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith('%%'):
            continue

        lower = line.lower()
        if lower.startswith('graph'):
            continue
        if lower.startswith('subgraph'):
            name = line[8:].strip()
            current_subgraph = name
            if name not in subgraph_order:
                subgraph_order.append(name)
            subgraph_nodes.setdefault(name, [])
            continue
        if lower == 'end':
            current_subgraph = None
            continue

        if '-->' in line:
            left, right = line.split('-->', 1)
            right_token = right.strip()
            source_part = left.strip()
            label = None
            if '--' in source_part:
                parts = source_part.split('--', 1)
                source_part = parts[0].strip()
                label = parts[1].strip()

            src_id, src_label, src_shape = _parse_node_token(source_part)
            dst_id, dst_label, dst_shape = _parse_node_token(right_token)
            _register_node(nodes, subgraph_nodes, standalone_nodes, src_id, src_label, src_shape, current_subgraph)
            _register_node(nodes, subgraph_nodes, standalone_nodes, dst_id, dst_label, dst_shape, current_subgraph)
            edges.append({'source': src_id, 'target': dst_id, 'label': label})
            continue

        node_id, node_label, shape = _parse_node_token(line)
        _register_node(nodes, subgraph_nodes, standalone_nodes, node_id, node_label, shape, current_subgraph)

    return {
        'nodes': nodes,
        'edges': edges,
        'subgraph_nodes': subgraph_nodes,
        'subgraph_order': subgraph_order,
        'standalone_nodes': standalone_nodes,
    }


def _draw_mermaid_structure(structure: dict, output_path: Path) -> bool:
    try:
        import numpy as np  # type: ignore
    except Exception as exc:  # noqa: BLE001
        print(f"[AVISO] Falha ao importar numpy para renderizar diagrama: {exc}")
        return False

    columns: List[Tuple[str, List[str]]] = []
    for name in structure['subgraph_order']:
        nodes = structure['subgraph_nodes'].get(name, [])
        if nodes:
            columns.append((name, nodes))

    standalone = structure['standalone_nodes']
    if standalone:
        columns.append(("Outros", standalone))

    if not columns:
        columns.append(("Fluxo", list(structure['nodes'].keys())))

    column_count = len(columns)
    max_nodes = max((len(nodes) for _, nodes in columns), default=1)

    # Ajustes de layout: elementos menores para caber em A4
    box_w = 220
    box_h = 90
    column_spacing = box_w + 140
    row_spacing = box_h + 50
    margin_x = 140
    margin_y = 140
    header_gap = 60

    width = int(margin_x * 2 + box_w + column_spacing * max(0, column_count - 1))
    height = int(margin_y * 2 + box_h + row_spacing * max(0, max_nodes - 1) + header_gap)

    max_width = 1550
    max_height = 1000
    scale_factor = min(1.0, max_width / max(width, 1), max_height / max(height, 1))

    if scale_factor < 1.0:
        box_w = int(box_w * scale_factor)
        box_h = int(box_h * scale_factor)
        column_spacing = int(column_spacing * scale_factor)
        row_spacing = int(row_spacing * scale_factor)
        margin_x = int(margin_x * scale_factor)
        margin_y = int(margin_y * scale_factor)
        header_gap = int(header_gap * scale_factor)
        width = int(width * scale_factor)
        height = int(height * scale_factor)
        font_scale_node = 0.55 * scale_factor + 0.35
        font_scale_title = 0.7 * scale_factor + 0.25
        font_scale_edge = 0.5 * scale_factor + 0.25
    else:
        font_scale_node = 0.6
        font_scale_title = 0.7
        font_scale_edge = 0.5

    img = np.full((height, width, 3), 255, dtype='uint8')

    positions: dict = {}

    font = cv2.FONT_HERSHEY_SIMPLEX
    node_color = (57, 89, 153)
    node_fill = (240, 246, 255)
    decision_color = (180, 130, 45)
    edge_color = (90, 90, 90)

    for col_idx, (title, node_ids) in enumerate(columns):
        center_x = int(margin_x + col_idx * column_spacing)
        if node_ids:
            title_size, _ = cv2.getTextSize(title, font, font_scale_title, 2)
            title_x = center_x - title_size[0] // 2
            title_y = int(margin_y - header_gap)
            cv2.putText(img, title, (title_x, title_y), font, font_scale_title, (60, 60, 60), 2, cv2.LINE_AA)

        for row_idx, node_id in enumerate(node_ids):
            center_y = int(margin_y + row_idx * row_spacing)
            positions[node_id] = (center_x, center_y)

    # Garante que todos os nós tenham posição
    for node_id in structure['nodes']:
        if node_id not in positions:
            col_idx = 0
            center_x = int(margin_x + col_idx * column_spacing)
            center_y = int(margin_y + len(positions) * row_spacing)
            positions[node_id] = (center_x, center_y)

    # Desenha nós
    for node_id, info in structure['nodes'].items():
        cx, cy = positions[node_id]
        label = info.get('label', node_id)
        shape = info.get('shape', 'rect')

        if shape == 'diamond':
            half_w = box_w // 2
            half_h = box_h // 2
            points = np.array([
                (cx, cy - half_h),
                (cx + half_w, cy),
                (cx, cy + half_h),
                (cx - half_w, cy),
            ], dtype=np.int32)
            cv2.fillPoly(img, [points], node_fill)
            cv2.polylines(img, [points], True, decision_color, 3, cv2.LINE_AA)
            text_color = (60, 60, 60)
        elif shape == 'circle':
            radius = min(box_w, box_h) // 2
            cv2.circle(img, (cx, cy), radius, node_fill, thickness=-1, lineType=cv2.LINE_AA)
            cv2.circle(img, (cx, cy), radius, node_color, thickness=3, lineType=cv2.LINE_AA)
            text_color = (60, 60, 60)
        else:
            top_left = (int(cx - box_w / 2), int(cy - box_h / 2))
            bottom_right = (int(cx + box_w / 2), int(cy + box_h / 2))
            cv2.rectangle(img, top_left, bottom_right, node_fill, thickness=-1)
            cv2.rectangle(img, top_left, bottom_right, node_color, thickness=3, lineType=cv2.LINE_AA)
            text_color = (60, 60, 60)

        max_chars = max(16, int(26 * scale_factor)) if scale_factor < 1.0 else 26
        wrapped = textwrap.wrap(label, width=max_chars) or [label]
        line_height = int(28 * scale_factor + 8)
        total_text_height = len(wrapped) * line_height
        text_start_y = cy - total_text_height // 2 + 10
        for idx, line in enumerate(wrapped):
            line = line.replace('"', '"')
            size, _ = cv2.getTextSize(line, font, font_scale_node, 2)
            text_x = cx - size[0] // 2
            text_y = text_start_y + idx * line_height
            cv2.putText(img, line, (text_x, text_y), font, font_scale_node, text_color, 2, cv2.LINE_AA)

    def _arrow_points(src, dst):
        sx, sy = positions[src]
        dx, dy = positions[dst]
        diff_x = dx - sx
        diff_y = dy - sy
        if abs(diff_x) >= abs(diff_y):
            direction = 1 if diff_x >= 0 else -1
            start = (int(sx + direction * box_w / 2), sy)
            end = (int(dx - direction * box_w / 2), dy)
        else:
            direction = 1 if diff_y >= 0 else -1
            start = (sx, int(sy + direction * box_h / 2))
            end = (dx, int(dy - direction * box_h / 2))
        return start, end

    for edge in structure['edges']:
        source = edge['source']
        target = edge['target']
        if source not in positions or target not in positions:
            continue
        start, end = _arrow_points(source, target)
        cv2.arrowedLine(img, start, end, edge_color, 2, cv2.LINE_AA, tipLength=0.02)

        label = edge.get('label')
        if label:
            midpoint = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
            size, _ = cv2.getTextSize(label, font, font_scale_edge, 1)
            label_pos = (midpoint[0] - size[0] // 2, midpoint[1] - 10)
            cv2.putText(img, label, label_pos, font, font_scale_edge, edge_color, 1, cv2.LINE_AA)

    return _write_image_bgr(img, output_path)


def _render_mermaid_locally(code: str, output_path: Path) -> bool:
    try:
        structure = _parse_mermaid_structure(code)
        if not structure['nodes']:
            return False
        return _draw_mermaid_structure(structure, output_path)
    except Exception as error:  # noqa: BLE001
        print(f"[AVISO] Falha ao renderizar diagrama Mermaid localmente: {error}")
        return False


def build_docx(
    soup: BeautifulSoup,
    metadata: dict,
    logo_path: Optional[Path],
    separator_path: Optional[Path],
    footer_banner_path: Optional[Path],
    output_path: Optional[Path] = None,
) -> Optional[bytes]:
    try:
        doc = Document()
    except Exception as error:
        print(f"[AVISO] Falha ao inicializar documento DOCX: {error}")
        return False

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(18)

    _configure_header(section, metadata, logo_path, separator_path)
    _configure_footer(section, metadata, footer_banner_path)

    if doc.paragraphs:
        first_para = doc.paragraphs[0]._element
        first_para.getparent().remove(first_para)

    content_width = section.page_width - section.left_margin - section.right_margin
    content_width_inches = content_width / 914400

    _add_html_content(doc, list(soup.children), content_width_inches)

    from io import BytesIO
    buf = BytesIO()
    try:
        doc.save(buf)
        docx_bytes = buf.getvalue()
    except Exception as error:
        print(f"[AVISO] Falha ao gerar DOCX: {error}")
        return None
    
    if output_path:
        try:
            output_path.write_bytes(docx_bytes)
        except Exception as error:
            print(f"[AVISO] Falha ao salvar DOCX: {error}")
    
    return docx_bytes


def _add_field(paragraph, field_code: str, font_size: float = 10) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    run._r.append(fld_char)

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = field_code
    run._r.append(instr_text)

    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'separate')
    run._r.append(fld_char)

    result_run = paragraph.add_run()
    result_run.text = "0"
    result_run.font.name = 'Calibri'
    if font_size:
        result_run.font.size = Pt(font_size)

    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'end')
    result_run._r.append(fld_char)


def _add_page_counter(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(paragraph, 'PAGE')
    slash_run = paragraph.add_run('/')
    slash_run.font.name = 'Calibri'
    slash_run.font.size = Pt(10)
    _add_field(paragraph, 'NUMPAGES')


def _resolve_asset_path(src: str) -> Optional[Path]:
    if not src:
        return None
    candidate = Path(src)
    if not candidate.is_absolute():
        candidate = (OUT_DIR / src).resolve()
    if candidate.exists():
        return candidate
    return None


def _parse_length_from_attr(value: Optional[str]) -> Optional[float]:
    if not value:
        return None
    text = str(value).strip().lower()
    if not text:
        return None
    multiplier = 1.0
    if text.endswith("px"):
        text = text[:-2].strip()
        multiplier = 1.0 / 96.0
    elif text.endswith("pt"):
        text = text[:-2].strip()
        multiplier = 1.0 / 72.0
    elif text.endswith("cm"):
        text = text[:-2].strip()
        multiplier = 1.0 / 2.54
    elif text.endswith("mm"):
        text = text[:-2].strip()
        multiplier = 1.0 / 25.4
    elif text.endswith("in"):
        text = text[:-2].strip()
        multiplier = 1.0
    try:
        numeric = float(text)
    except (TypeError, ValueError):
        return None
    inches = numeric * multiplier
    if inches <= 0:
        return None
    return inches


def _get_image_dimensions(img_path: Path) -> Optional[Tuple[int, int]]:
    try:
        img = cv2.imread(str(img_path))
        if img is not None:
            height, width = img.shape[:2]
            return width, height
        # Fallback: parse dimensions from file headers to handle Unicode path issues
        data = img_path.read_bytes()
        # PNG header
        if len(data) >= 24 and data[:8] == b'\x89PNG\r\n\x1a\n':
            # IHDR chunk starts at byte 8+8
            width = int.from_bytes(data[16:20], 'big')
            height = int.from_bytes(data[20:24], 'big')
            if width > 0 and height > 0:
                return width, height
        # JPEG header (scan for SOF0/2 markers)
        i = 0
        while i + 9 < len(data) and data[i] == 0xFF:
            if data[i + 1] in (0xC0, 0xC2):
                # SOF marker length is at i+2..i+3
                h = int.from_bytes(data[i + 5:i + 7], 'big')
                w = int.from_bytes(data[i + 7:i + 9], 'big')
                if w > 0 and h > 0:
                    return w, h
                break
            if data[i + 1] == 0xDA:  # SOS
                break
            seg_len = int.from_bytes(data[i + 2:i + 4], 'big')
            if seg_len <= 0:
                break
            i += 2 + seg_len
    except Exception as error:  # noqa: BLE001
        print(f"[AVISO] Não foi possível ler dimensões da imagem {img_path}: {error}")
    return None


def _add_image_to_paragraph(paragraph, img_tag: Tag, content_width_inches: float) -> None:
    src = img_tag.get('src', '')
    img_path = _resolve_asset_path(src)
    if not img_path:
        print(f"[AVISO] Imagem não encontrada para inserção no DOCX: {src}")
        return

    width_inches = _parse_length_from_attr(img_tag.get('width'))
    height_inches = _parse_length_from_attr(img_tag.get('height'))

    max_width_inches = max(content_width_inches - 0.4, 1.0)
    resolved_path = None
    try:
        resolved_path = img_path.resolve()
    except Exception:
        resolved_path = img_path

    is_diagram = False
    is_print = False
    try:
        diagram_root = DIAGRAMS_DIR.resolve()
        is_diagram = resolved_path and diagram_root in resolved_path.parents
    except Exception:
        is_diagram = False

    try:
        prints_root = ASSETS_DIR.resolve()
        is_print = resolved_path and prints_root in resolved_path.parents
    except Exception:
        is_print = False

    figure_classes = img_tag.get('class', [])
    if isinstance(figure_classes, (list, tuple, set)) and 'video-print' in figure_classes:
        is_print = True

    if width_inches is None and height_inches is None:
        width_inches = max_width_inches
    elif width_inches is not None and width_inches > max_width_inches:
        width_inches = max_width_inches

    if is_diagram:
        dimensions = _get_image_dimensions(img_path)
        if dimensions:
            fig_width_px, fig_height_px = dimensions
            if fig_width_px > 0:
                aspect = fig_height_px / fig_width_px
                target_width = min(max_width_inches, MAX_DIAGRAM_WIDTH_INCHES)
                target_height = target_width * aspect
                if target_height > MAX_DIAGRAM_HEIGHT_INCHES and aspect > 0:
                    target_height = MAX_DIAGRAM_HEIGHT_INCHES
                    target_width = target_height / aspect
                width_inches = target_width
                height_inches = target_height
    else:
        scale_factor = PRINT_IMAGE_SCALE if is_print else 1.0
        if width_inches is not None:
            width_inches = max(width_inches * scale_factor, 0.1)
        if height_inches is not None:
            height_inches = max(height_inches * scale_factor, 0.1)
        if width_inches is None and height_inches is None:
            width_inches = max_width_inches * scale_factor

    run = paragraph.add_run()
    try:
        if width_inches is not None and height_inches is not None:
            run.add_picture(str(img_path), width=Inches(width_inches), height=Inches(height_inches))
        elif width_inches is not None:
            run.add_picture(str(img_path), width=Inches(width_inches))
        elif height_inches is not None:
            run.add_picture(str(img_path), height=Inches(height_inches))
        else:
            run.add_picture(str(img_path), width=Inches(max_width_inches))
    except Exception as error:
        print(f"[AVISO] Falha ao inserir imagem no DOCX ({img_path}): {error}")
        paragraph._element.remove(run._r)


def _add_runs(doc: Document, paragraph, node, content_width_inches: float, bold: bool = False, italic: bool = False) -> None:
    if isinstance(node, NavigableString):
        text = str(node)
        if not text or not text.strip():
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        return

    if not isinstance(node, Tag):
        return

    name = (node.name or '').lower()
    if name == 'br':
        paragraph.add_run().add_break()
        return

    if name == 'figure':
        _add_figure(doc, node, content_width_inches)
        return

    if name == 'img':
        _add_image_to_paragraph(paragraph, node, content_width_inches)
        return

    if name in {'ul', 'ol'}:
        _add_list(doc, node, ordered=(name == 'ol'), content_width_inches=content_width_inches)
        return

    next_bold = bold or name in {'strong', 'b'}
    next_italic = italic or name in {'em', 'i'}

    for child in node.children:
        _add_runs(doc, paragraph, child, content_width_inches, next_bold, next_italic)


def _add_list(doc: Document, list_tag: Tag, ordered: bool, content_width_inches: float, level: int = 0) -> None:
    style_name = 'List Number' if ordered else 'List Bullet'
    base_indent = 28.35
    indent_step = 18

    for item in list_tag.find_all('li', recursive=False):
        meaningful_children = [child for child in item.children if not (isinstance(child, NavigableString) and not str(child).strip())]
        if len(meaningful_children) == 1 and isinstance(meaningful_children[0], Tag) and (meaningful_children[0].name or '').lower() == 'figure':
            _add_figure(doc, meaningful_children[0], content_width_inches)
            continue

        paragraph = doc.add_paragraph(style=style_name)
        paragraph.paragraph_format.left_indent = Pt(base_indent + level * indent_step)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.space_before = Pt(3)
        paragraph.paragraph_format.space_after = Pt(6 if level == 0 else 3)
        paragraph.paragraph_format.line_spacing = 1.15

        for child in item.children:
            if isinstance(child, NavigableString):
                _add_runs(doc, paragraph, child, content_width_inches)
                continue

            if not isinstance(child, Tag):
                continue

            child_name = (child.name or '').lower()
            if child_name in {'ul', 'ol'}:
                _add_list(doc, child, ordered=(child_name == 'ol'), content_width_inches=content_width_inches, level=level + 1)
            else:
                _add_runs(doc, paragraph, child, content_width_inches)

        if not paragraph.text.strip():
            parent = paragraph._element.getparent()
            if parent is not None:
                parent.remove(paragraph._element)

    if level == 0:
        doc.add_paragraph()


def _add_figure(doc: Document, figure_tag: Tag, content_width_inches: float) -> None:
    img = figure_tag.find('img')
    if not img:
        return
    src = img.get('src', '')
    img_path = _resolve_asset_path(src)
    if not img_path:
        print(f"[AVISO] Figura não encontrada para inserção no DOCX: {src}")
        return

    paragraph = doc.add_paragraph()
    paragraph.style = doc.styles['Normal']
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.left_indent = Pt(0)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    width_inches = _parse_length_from_attr(img.get('width'))
    height_inches = _parse_length_from_attr(img.get('height'))
    max_width_inches = max(content_width_inches - 0.4, 1.0)
    resolved_path = None
    try:
        resolved_path = img_path.resolve()
    except Exception:
        resolved_path = img_path

    classes = figure_tag.get('class', []) or []
    if isinstance(classes, str):
        classes = [classes]

    is_diagram = False
    try:
        diagram_root = DIAGRAMS_DIR.resolve()
        is_diagram = resolved_path and diagram_root in resolved_path.parents
    except Exception:
        is_diagram = False
    if 'process-diagram' in classes:
        is_diagram = True

    is_print = False
    try:
        prints_root = ASSETS_DIR.resolve()
        is_print = resolved_path and prints_root in resolved_path.parents
    except Exception:
        is_print = False
    if 'video-print' in classes:
        is_print = True

    if width_inches is None and height_inches is None:
        width_inches = max_width_inches
    elif width_inches is not None and width_inches > max_width_inches:
        width_inches = max_width_inches

    if is_diagram:
        dimensions = _get_image_dimensions(img_path)
        if dimensions:
            fig_width_px, fig_height_px = dimensions
            if fig_width_px > 0:
                aspect = fig_height_px / fig_width_px
                target_width = min(max_width_inches, MAX_DIAGRAM_WIDTH_INCHES)
                target_height = target_width * aspect
                if target_height > MAX_DIAGRAM_HEIGHT_INCHES and aspect > 0:
                    target_height = MAX_DIAGRAM_HEIGHT_INCHES
                    target_width = target_height / aspect
                width_inches = target_width
                height_inches = target_height
    else:
        scale_factor = PRINT_IMAGE_SCALE if is_print else 1.0
        if width_inches is not None:
            width_inches = max(width_inches * scale_factor, 0.1)
        if height_inches is not None:
            height_inches = max(height_inches * scale_factor, 0.1)
        if width_inches is None and height_inches is None:
            width_inches = max_width_inches * scale_factor

    try:
        if width_inches is not None and height_inches is not None:
            run.add_picture(str(img_path), width=Inches(width_inches), height=Inches(height_inches))
        elif width_inches is not None:
            run.add_picture(str(img_path), width=Inches(width_inches))
        elif height_inches is not None:
            run.add_picture(str(img_path), height=Inches(height_inches))
        else:
            run.add_picture(str(img_path), width=Inches(max_width_inches))
    except Exception as error:
        paragraph.text = ''
        print(f"[AVISO] Falha ao inserir figura no DOCX ({img_path}): {error}")
        return

    caption = figure_tag.find('figcaption')
    if caption:
        caption_text = caption.get_text(strip=True)
        if not hasattr(doc, '_figure_counter'):
            setattr(doc, '_figure_counter', 0)
        doc._figure_counter += 1  # type: ignore[attr-defined]
        caption_prefix = f"Figura {doc._figure_counter:02d} – "  # type: ignore[attr-defined]

        caption_para = doc.add_paragraph()
        try:
            caption_para.style = doc.styles['Caption']
        except KeyError:
            caption_para.style = doc.styles['Normal']
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.paragraph_format.left_indent = Pt(0)
        caption_para.paragraph_format.first_line_indent = Pt(0)
        caption_para.paragraph_format.space_before = Pt(1)
        caption_para.paragraph_format.space_after = Pt(6)

        caption_run_label = caption_para.add_run(caption_prefix)
        caption_run_label.bold = True
        caption_run_label.font.name = 'Calibri'
        caption_run_label.font.size = Pt(9)

        caption_run_text = caption_para.add_run(caption_text)
        caption_run_text.italic = True
        caption_run_text.font.name = 'Calibri'
        caption_run_text.font.size = Pt(9)


def _add_table_from_tag(doc: Document, table_tag: Tag, content_width_inches: float) -> None:
    rows = table_tag.find_all('tr')
    if not rows:
        return

    parsed_rows = []
    max_cols = 0
    for row in rows:
        cells = row.find_all(['td', 'th'])
        parsed_rows.append(cells)
        max_cols = max(max_cols, len(cells))

    docx_table = doc.add_table(rows=len(parsed_rows), cols=max_cols)
    docx_table.style = 'Table Grid'
    docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, cells in enumerate(parsed_rows):
        for c_idx in range(max_cols):
            cell = docx_table.cell(r_idx, c_idx)
            paragraph = cell.paragraphs[0]
            paragraph.text = ''
            if c_idx < len(cells):
                content = cells[c_idx]
                is_header = (content.name or '').lower() == 'th'
                _add_runs(doc, paragraph, content, content_width_inches, bold=is_header)

    doc.add_paragraph()


def _add_html_content(doc: Document, nodes, content_width_inches: float) -> None:
    for child in nodes:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                paragraph = doc.add_paragraph(text)
                paragraph.paragraph_format.space_after = Pt(6)
            continue

        if not isinstance(child, Tag):
            continue

        name = (child.name or '').lower()

        if name == 'p':
            paragraph = doc.add_paragraph()
            _add_runs(doc, paragraph, child, content_width_inches)
            paragraph.paragraph_format.space_after = Pt(6)
            continue

        if name == 'h2':
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(child.get_text(strip=True).upper())
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            continue

        if name == 'h3':
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(child.get_text(strip=True))
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            continue

        if name == 'ul':
            _add_list(doc, child, ordered=False, content_width_inches=content_width_inches)
            continue

        if name == 'ol':
            _add_list(doc, child, ordered=True, content_width_inches=content_width_inches)
            continue

        if name == 'figure':
            _add_figure(doc, child, content_width_inches)
            continue

        if name == 'table':
            _add_table_from_tag(doc, child, content_width_inches)
            continue

        if name == 'blockquote':
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Pt(18)
            paragraph.paragraph_format.space_after = Pt(6)
            _add_runs(doc, paragraph, child, content_width_inches)
            continue

        if name == 'hr':
            doc.add_page_break()
            continue

        _add_html_content(doc, child.children, content_width_inches)


def _configure_header(section, metadata: dict, logo_path: Optional[Path], separator_path: Optional[Path]) -> None:
    header = section.header
    hdr_element = header._element
    for child in list(hdr_element):
        hdr_element.remove(child)

    usable_width = section.page_width - section.left_margin - section.right_margin
    table = header.add_table(rows=2, cols=4, width=usable_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = None

    top_row = table.rows[0]
    top_row.height = Pt(59.55)
    top_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
    bottom_row = table.rows[1]
    bottom_row.height = Pt(15)
    bottom_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    column_widths_pt = [21.05, 80.0, 296.0, 175.0]
    for row in table.rows:
        for idx, width_pt in enumerate(column_widths_pt):
            width = Inches(width_pt / 72)
            row.cells[idx].width = width
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    title_text = f"{metadata.get('doc_type', '')} – {metadata.get('doc_title', '')}".upper()

    # Coluna à esquerda permanece vazia para manter o layout
    empty_cell = table.cell(0, 0)
    empty_cell.text = ''
    empty_para = empty_cell.paragraphs[0]
    empty_para.paragraph_format.space_before = Pt(0)
    empty_para.paragraph_format.space_after = Pt(0)

    # Logo
    logo_cell = table.cell(0, 1)
    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_para.paragraph_format.space_before = Pt(0)
    logo_para.paragraph_format.space_after = Pt(0)
    if logo_path and Path(logo_path).exists():
        try:
            logo_para.add_run().add_picture(str(logo_path), width=Inches(111 / 96))
        except Exception as error:
            print(f"[AVISO] Falha ao inserir logo no cabeçalho: {error}")

    # Title
    title_cell = table.cell(0, 2)
    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.left_indent = Pt(24.85)
    title_para.paragraph_format.line_spacing = 1.08
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    print(f"[DEBUG] Title: {repr(title_text)}", file=sys.stderr)
    title_run = title_para.add_run(title_text)
    title_run.bold = True
    title_run.font.name = 'Calibri'
    title_run.font.size = Pt(14)

    # Meta data
    meta_cell = table.cell(0, 3)
    for paragraph in list(meta_cell.paragraphs):
        p_elem = paragraph._element
        p_elem.getparent().remove(p_elem)

    meta_entries = [
        ("Código:", metadata.get('doc_code', '')),
        ("Data de Emissão:", metadata.get('doc_issue', '')),
        ("Revisão:", metadata.get('doc_revision', '')),
    ]
    for label, value in meta_entries:
        para = meta_cell.add_paragraph()
        para.paragraph_format.left_indent = Pt(10.5)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run_label = para.add_run(label + ' ')
        run_label.bold = True
        run_label.font.name = 'Calibri'
        run_label.font.size = Pt(10)
        run_value = para.add_run(value)
        run_value.font.name = 'Calibri'
        run_value.font.size = Pt(10)

    # Separator row
    separator_row = table.rows[1]
    merged_cell = separator_row.cells[0]
    for idx in range(1, 4):
        merged_cell = merged_cell.merge(separator_row.cells[idx])
    separator_para = merged_cell.paragraphs[0]
    separator_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    separator_para.paragraph_format.space_before = Pt(0)
    separator_para.paragraph_format.space_after = Pt(0)
    if separator_path and Path(separator_path).exists():
        try:
            # Converter largura utilizável para polegadas
            usable_width_inches = (usable_width / 914400)
            # Altura máxima desejada para a faixa separadora (aplainar se for muito alta)
            max_sep_height_inches = 0.35  # ~9 mm
            dims = _get_image_dimensions(Path(separator_path))
            if dims and dims[0] > 0:
                w_px, h_px = dims
                # Altura natural quando ajustamos a largura para a largura utilizável (assumindo 96 DPI)
                natural_h_inches = (h_px / 96.0) * (usable_width_inches / max(w_px / 96.0, 1e-6))
                if natural_h_inches > max_sep_height_inches:
                    # Insere e força altura reduzida (pode distorcer, conforme solicitado para "esmagar")
                    pic = separator_para.add_run().add_picture(str(separator_path), width=Inches(usable_width_inches))
                    try:
                        pic.height = Inches(max_sep_height_inches)
                    except Exception:
                        pass
                else:
                    separator_para.add_run().add_picture(str(separator_path), width=Inches(usable_width_inches))
            else:
                # Sem dimensões detectáveis: força largura e altura máxima
                pic = separator_para.add_run().add_picture(str(separator_path), width=Inches(usable_width_inches))
                try:
                    pic.height = Inches(max_sep_height_inches)
                except Exception:
                    pass
        except Exception as error:
            print(f"[AVISO] Falha ao inserir faixa separadora no cabeçalho: {error}")
    else:
        print(f"[AVISO] Separador não encontrado em: {separator_path}")


def _configure_footer(section, metadata: dict, footer_banner_path: Optional[Path]) -> None:
    footer = section.footer
    ftr_element = footer._element
    for child in list(ftr_element):
        ftr_element.remove(child)

    banner_para = footer.add_paragraph()
    banner_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    banner_para.paragraph_format.space_before = Pt(0)
    banner_para.paragraph_format.space_after = Pt(6)
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_width_inches = usable_width / 914400
    if footer_banner_path and Path(footer_banner_path).exists():
        try:
            banner_para.add_run().add_picture(str(footer_banner_path), width=usable_width)
        except Exception as error:
            print(f"[AVISO] Falha ao inserir faixa no rodapé: {error}")

    table = footer.add_table(rows=2, cols=3, width=usable_width)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = None

    column_width_inches = [usable_width_inches / 3.0] * 3
    for row in table.rows:
        for idx, width_inches in enumerate(column_width_inches):
            row.cells[idx].width = Inches(width_inches)
            row.cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    labels = ["Elaboração/ Revisão:", "Aprovação:", "Páginas:"]
    values = [metadata.get('elaboracao', '—') or '—', metadata.get('aprovacao', '—') or '—']

    for idx, label in enumerate(labels):
        cell = table.cell(0, idx)
        paragraph = cell.paragraphs[0]
        paragraph.text = ''
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(label)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    value_cells = [table.cell(1, 0), table.cell(1, 1)]
    for cell, value in zip(value_cells, values):
        paragraph = cell.paragraphs[0]
        paragraph.text = ''
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(value)
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pages_paragraph = table.cell(1, 2).paragraphs[0]
    pages_paragraph.text = ''
    pages_paragraph.paragraph_format.space_before = Pt(0)
    pages_paragraph.paragraph_format.space_after = Pt(0)
    _add_page_counter(pages_paragraph)


def main(layout_assets_dir: Optional[Path] = None) -> int:
    ensure_dirs()
    # Redireciona stdout textual para stderr para evitar poluir o fluxo binário do DOCX
    orig_stdout = sys.stdout
    orig_out_buffer = sys.stdout.buffer
    sys.stdout = sys.stderr
    # Ler md de stdin como bytes e decodear
    md_bytes = sys.stdin.buffer.read()
    md_text = md_bytes.decode('utf-8', errors='replace')  # Substituir caracteres inválidos
    md_text = clean_text(md_text)
    if not md_text.strip():
        print("[ERRO] Nenhum conteúdo MD fornecido via stdin")
        return 1
    
    # Vídeo é opcional - se não existir, prints não serão extraídos
    video_available = VIDEO_FILE.exists() if _VIDEO_ENV else False

    md_processed, occurrences = replace_print_placeholders(md_text)
    md_processed, mermaid_paths = replace_mermaid_blocks(md_processed)

    # Extrai frames apenas se vídeo disponível
    if video_available:
        for occurrence in occurrences:
            seconds = parse_timestamp_to_seconds(occurrence.timestamp)
            ok = extract_frame(VIDEO_FILE, seconds, occurrence.image_path, coords=occurrence.coords)
            if not ok:
                print(f"[AVISO] Falha ao capturar {occurrence.timestamp}. Placeholder gerado.")
    elif occurrences:
        print(f"[AVISO] Vídeo não disponível. {len(occurrences)} placeholder(s) de print não serão preenchidos com imagens.")

    if mermaid_paths:
        print(f"[OK] {len(mermaid_paths)} diagrama(s) Mermaid disponível(is) em: {DIAGRAMS_DIR}")

    # Converte markdown para HTML intermediário e monta BeautifulSoup para processar no DOCX
    html_body_raw = markdown(md_processed, extensions=['extra', 'tables', 'sane_lists', 'toc'])
    soup = BeautifulSoup(html_body_raw, 'html.parser')
    inject_print_figures(soup, occurrences)

    first_h1 = soup.find('h1')
    title_text = "Documento"
    if first_h1 and first_h1.text.strip():
        title_text = clean_text(first_h1.text.strip())
        first_h1.decompose()

    metadata = load_metadata(title_text)

    # Ler diretório de assets customizados somente via variável de ambiente (não usar pasta do projeto)
    if layout_assets_dir is None:
        env_dir = os.environ.get("LAYOUT_ASSETS_DIR")
        if env_dir:
            p = Path(env_dir)
            if p.exists():
                layout_assets_dir = p

    logo_path = find_logo(layout_assets_dir)
    separator_path = find_model_separator(layout_assets_dir)
    footer_banner_path = find_model_footer_banner(layout_assets_dir)

    print(f"[DEBUG] Layout assets dir: {layout_assets_dir}")
    print(f"[DEBUG] Logo: {logo_path}")
    print(f"[DEBUG] Separator: {separator_path}")
    print(f"[DEBUG] Footer banner: {footer_banner_path}")

    docx_bytes = build_docx(
        soup,
        metadata=metadata,
        logo_path=logo_path,
        separator_path=separator_path,
        footer_banner_path=footer_banner_path,
        output_path=None,  # Não salvar em arquivo
    )
    if docx_bytes is None:
        print('[AVISO] Falha ao gerar DOCX automaticamente.')
        return 1

    # Escrever bytes puros no stdout original (sem prints misturados)
    try:
        orig_out_buffer.write(docx_bytes)
        orig_out_buffer.flush()
    except Exception as exc:
        print(f"[ERRO] Falha ao escrever DOCX em stdout: {exc}")
        return 1
    print(f"[OK] Documento DOCX gerado ({len(docx_bytes)} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
